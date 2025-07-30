from PyPDF2 import PdfReader
from PyPDF2.errors import PdfReadError
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
import pdfplumber
import unicodedata
import logging
from tempfile import TemporaryDirectory
import os
import re
from .legacy_kannada import (
    convert_legacy_to_unicode,
    post_process_kannada_text,
    detect_legacy_encoding,
    validate_kannada_output,
    normalize_unicode,
    is_kannada_text,
)
from .hybrid_pdf_detector import detect_pdf_type, extract_images_from_pdf

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _detect_legacy_fonts_in_pdf(pdf_path: str) -> bool:
    """Detect if PDF contains legacy Kannada fonts that need conversion."""
    try:
        reader = PdfReader(pdf_path, strict=False)

        # Check font names in the PDF
        legacy_font_indicators = [
            "Nudi",
            "BRH",
            "Baraha",
            "KGP",
            "Kedage",
            "Malige",
            "Akshar",
            "Hubballi",
            "Mallige",
            "Sampige",
            "Tunga",
        ]

        for page in reader.pages:
            if "/Font" in page:
                fonts = page["/Font"]
                for font_ref in fonts.values():
                    if hasattr(font_ref, "get_object"):
                        font_obj = font_ref.get_object()
                        if "/BaseFont" in font_obj:
                            font_name = str(font_obj["/BaseFont"])
                            for indicator in legacy_font_indicators:
                                if indicator.lower() in font_name.lower():
                                    logger.info(f"Legacy font detected: {font_name}")
                                    return True

        return False

    except Exception as e:
        logger.warning(f"Could not analyze PDF fonts: {e}")
        return False


def _extract_text_with_encoding_detection(pdf_path: str) -> tuple[str, bool]:
    """Extract text and detect encoding issues."""
    extracted_text = ""
    has_legacy_fonts = False

    try:
        # First check for legacy fonts
        has_legacy_fonts = _detect_legacy_fonts_in_pdf(pdf_path)

        # Try multiple extraction methods
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                logger.debug(f"Extracting text from page {page_num}")

                # Primary extraction method
                page_text = page.extract_text() or ""

                # If no text, try with different layout parameters
                if not page_text.strip():
                    page_text = (
                        page.extract_text(layout=True, x_tolerance=3, y_tolerance=3)
                        or ""
                    )

                if page_text.strip():
                    extracted_text += page_text + "\n\n"
                else:
                    logger.warning(f"No text extracted from page {page_num}")

    except Exception as e:
        logger.error(f"Error extracting text: {e}")

    return extracted_text, has_legacy_fonts


def _process_extracted_text(text: str, has_legacy_fonts: bool = False) -> str:
    """Process extracted text with comprehensive Kannada handling."""
    if not text.strip():
        return ""

    logger.info(f"Processing text (legacy fonts: {has_legacy_fonts})")
    logger.info(f"Sample raw text: '{text[:100]}...'")

    # Step 1: Check if text is already good Unicode Kannada
    if is_kannada_text(text) and not has_legacy_fonts:
        logger.info("Text appears to be good Unicode Kannada - minimal processing")
        # Just normalize and clean lightly
        processed = normalize_unicode(text)
        processed = re.sub(r"\s+", " ", processed).strip()
        return processed

    # Step 2: Only apply legacy conversion if fonts detected or text is clearly corrupted
    is_legacy = has_legacy_fonts or detect_legacy_encoding(text)

    if is_legacy:
        logger.info("Applying legacy font conversion")
        text = convert_legacy_to_unicode(text)
        text = post_process_kannada_text(text, is_legacy=True)
    else:
        logger.info("Applying light processing only")
        # Light processing for digital PDFs
        text = normalize_unicode(text)
        text = re.sub(r"\s+", " ", text).strip()

    # Step 3: Final validation
    is_valid, issues = validate_kannada_output(text)
    if not is_valid:
        logger.warning(f"Text quality issues: {', '.join(issues)}")

        # Only try aggressive processing if text is really bad
        if not is_kannada_text(text) and len(text) > 20:
            logger.info("Attempting aggressive processing as fallback")
            text = convert_legacy_to_unicode(text)
            text = post_process_kannada_text(text, is_legacy=True)

    logger.info(f"Final processed text sample: '{text[:100]}...'")
    return text


def _set_kannada_font(run):
    """Set appropriate font for Kannada text rendering."""
    try:
        run.font.name = "Noto Sans Kannada"
        # Fallback fonts for better compatibility
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Kannada")
        run._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Kannada")
    except Exception as e:
        logger.warning(f"Could not set Kannada font: {e}")


def _add_page_header(document: Document, page_num: int, total_pages: int):
    """Add a formatted page header."""
    header = document.add_paragraph(f"ಪುಟ {page_num} / {total_pages}")
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.style = "Heading 2"

    # Set Kannada font for header
    for run in header.runs:
        _set_kannada_font(run)


def convert_pdf_to_word(
    input_pdf_path: str,
    output_docx_path: str,
    output_txt_path: str | None = None,
    *,
    title: str | None = None,
    author: str | None = None,
) -> tuple[str, str]:
    """Convert a digital PDF file to a Word document with enhanced Kannada support."""

    logger.info(f"Starting digital PDF conversion: {input_pdf_path}")

    # Initialize document
    document = Document()

    # Set output paths
    if output_txt_path is None:
        output_txt_path = os.path.splitext(output_docx_path)[0] + ".txt"

    # Set document metadata
    core = document.core_properties
    if title:
        core.title = title
    if author:
        core.author = author
    core.language = "kn-IN"  # Kannada locale

    # Validate PDF
    try:
        reader = PdfReader(input_pdf_path, strict=False)
        total_pages = len(reader.pages)
        logger.info(f"PDF has {total_pages} pages")
    except PdfReadError as exc:
        raise ValueError("Uploaded file is not a valid PDF.") from exc

    # Extract text with encoding detection
    try:
        raw_text, has_legacy_fonts = _extract_text_with_encoding_detection(
            input_pdf_path
        )

        if not raw_text.strip():
            logger.warning("No text extracted from PDF - might be scanned")
            # Add helpful message to document
            no_text_msg = (
                "ಈ PDF ಯಿಂದ ಯಾವುದೇ ಪಠ್ಯವನ್ನು ಹೊರತೆಗೆಯಲಾಗಲಿಲ್ಲ. "
                "ಇದು ಸ್ಕ್ಯಾನ್ ಮಾಡಿದ PDF ಆಗಿರಬಹುದು. "
                "'Scanned PDF' ಮೋಡ್ ಅನ್ನು ಪ್ರಯತ್ನಿಸಿ."
            )
            warning_para = document.add_paragraph(no_text_msg)
            for run in warning_para.runs:
                _set_kannada_font(run)
                run.bold = True

            full_text = no_text_msg
        else:
            # Process the extracted text
            processed_text = _process_extracted_text(raw_text, has_legacy_fonts)

            if processed_text.strip():
                # Add processed text to document
                text_paragraph = document.add_paragraph(processed_text)

                # Set Kannada font for all runs
                for run in text_paragraph.runs:
                    _set_kannada_font(run)

                full_text = processed_text
                logger.info("Text processing completed successfully")
            else:
                logger.error("Text processing returned empty result")
                error_msg = "ಪಠ್ಯ ಸಂಸ್ಕರಣೆಯಲ್ಲಿ ದೋಷ ಸಂಭವಿಸಿದೆ. " "PDF ಯಲ್ಲಿ ಪುರಾತನ ಫಾಂಟ್‌ಗಳಿರಬಹುದು."
                error_para = document.add_paragraph(error_msg)
                for run in error_para.runs:
                    _set_kannada_font(run)
                    run.italic = True

                full_text = error_msg

    except Exception as e:
        logger.error(f"Error during PDF processing: {e}")
        error_msg = f"PDF ಸಂಸ್ಕರಣೆಯಲ್ಲಿ ದೋಷ: {str(e)}"
        error_para = document.add_paragraph(error_msg)
        for run in error_para.runs:
            _set_kannada_font(run)
        full_text = error_msg

    # Save Word document
    try:
        document.save(output_docx_path)
        logger.info(f"Word document saved: {output_docx_path}")
    except Exception as e:
        logger.error(f"Failed to save Word document: {e}")
        raise

    # Save text file
    try:
        with open(output_txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        logger.info(f"Text file saved: {output_txt_path}")
    except Exception as e:
        logger.error(f"Failed to save text file: {e}")
        raise

    return output_docx_path, output_txt_path
