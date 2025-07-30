from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import qn
import pytesseract
from google.cloud import vision
from PIL import Image, ImageOps, ImageEnhance
from tempfile import TemporaryDirectory
import io
import os
import unicodedata
import cv2
import numpy as np
import logging
import time
import gc
from typing import Optional
from .legacy_kannada import (
    post_process_kannada_text,
    detect_legacy_encoding,
    validate_kannada_output,
    normalize_unicode,
    is_kannada_text,
)
from .kannada_image_preprocessor import preprocess_kannada_image

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def _preprocess_image_for_kannada(img):
    """Enhanced preprocessing specifically optimized for Kannada OCR using new pipeline."""
    logger.debug("Starting enhanced Kannada image preprocessing")

    try:
        # Use the new advanced preprocessing pipeline
        processed_array = preprocess_kannada_image(img)

        # Convert back to PIL Image for OCR compatibility
        processed_img = Image.fromarray(processed_array)

        logger.debug("Enhanced preprocessing completed successfully")
        return processed_img

    except Exception as e:
        logger.error(f"Enhanced preprocessing failed, using fallback: {e}")

        # Fallback to simple preprocessing if new pipeline fails
        try:
            # Convert PIL to OpenCV format
            opencv_img = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

            # Convert to grayscale
            gray = cv2.cvtColor(opencv_img, cv2.COLOR_BGR2GRAY)

            # Simple enhancement
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(gray)

            # Basic thresholding
            _, binary = cv2.threshold(
                enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
            )

            return Image.fromarray(binary)

        except Exception as e2:
            logger.error(f"Fallback preprocessing also failed: {e2}")
            # Return original image if all preprocessing fails
            return img


def _perform_tesseract_ocr(
    image, language: str = "kan", timeout: Optional[int] = 30
) -> str:
    """Perform Tesseract OCR with Kannada-specific configuration."""
    try:
        logger.debug("Starting Tesseract OCR")
        start_time = time.time()

        # Enhanced Tesseract configuration for Kannada magazines
        custom_config = r"--oem 3 --psm 6 -c tessedit_char_whitelist=ಅ-ೞ೦-೯ -c preserve_interword_spaces=1"

        # First attempt with Kannada whitelist and preserve spaces
        text = pytesseract.image_to_string(
            image, lang=language, config=custom_config, timeout=timeout
        )

        elapsed = time.time() - start_time
        logger.debug(
            f"Tesseract OCR completed in {elapsed:.2f}s, extracted {len(text)} characters"
        )

        # If no text found, try alternative configurations
        if not text.strip():
            logger.warning("No text with whitelist, trying without restrictions")
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang=language,
                    config=r"--oem 3 --psm 6 -c preserve_interword_spaces=1",
                    timeout=timeout // 2,
                )
            except Exception as e:
                logger.warning(f"Fallback OCR failed: {e}")

        # If still no text, try different PSM modes
        if not text.strip():
            logger.warning("Trying alternative PSM modes")
            for psm in [3, 4, 8]:
                try:
                    text = pytesseract.image_to_string(
                        image,
                        lang=language,
                        config=f"--oem 3 --psm {psm} -c preserve_interword_spaces=1",
                        timeout=15,
                    )
                    if text.strip():
                        logger.info(f"Success with PSM mode {psm}")
                        break
                except Exception:
                    continue

        return text

    except pytesseract.TesseractError as e:
        if "timeout" in str(e).lower():
            logger.error(f"Tesseract OCR timed out after {timeout} seconds")
            raise TimeoutError(
                f"Tesseract OCR timed out after {timeout} seconds"
            ) from e
        logger.error(f"Tesseract OCR error: {e}")
        return ""
    except Exception as e:
        logger.error(f"Tesseract OCR failed: {e}")
        return ""


def _perform_google_vision_ocr(
    image, language: str = "kan", timeout: Optional[int] = 30
) -> str:
    """Perform Google Vision OCR with enhanced Kannada optimization."""
    try:
        logger.debug("Starting Google Vision OCR")
        start_time = time.time()

        client = vision.ImageAnnotatorClient()

        # Convert image to bytes with optimized compression
        buffer = io.BytesIO()
        image.save(
            buffer, format="PNG", optimize=False, compress_level=1
        )  # Fast compression
        buffer.seek(0)

        # Create Vision API image object
        vision_image = vision.Image(content=buffer.getvalue())

        # Enhanced image context for better Kannada recognition
        image_context = vision.ImageContext(
            language_hints=["kn", "en"],  # Use proper language codes
            text_detection_params=vision.TextDetectionParams(
                enable_text_detection_confidence_score=True,
                advanced_ocr_options=[
                    vision.TextDetectionParams.AdvancedOcrOption.LEGACY_LAYOUT
                ],
            ),
        )

        # Use document_text_detection for better layout preservation
        response = client.document_text_detection(
            image=vision_image, image_context=image_context, timeout=timeout
        )

        if response.error.message:
            raise Exception(f"Google Vision API error: {response.error.message}")

        text = (
            response.full_text_annotation.text if response.full_text_annotation else ""
        )

        elapsed = time.time() - start_time
        logger.debug(
            f"Google Vision OCR completed in {elapsed:.2f}s, extracted {len(text)} characters"
        )

        # Log sample of extracted text for debugging
        if text:
            sample = text[:100].replace("\n", " ")
            logger.debug(f"Sample Google Vision text: {sample}")

        return text

    except Exception as e:
        if "deadline exceeded" in str(e).lower() or "timeout" in str(e).lower():
            logger.error(f"Google Vision OCR timed out after {timeout} seconds")
            raise TimeoutError(
                f"Google Vision OCR timed out after {timeout} seconds"
            ) from e
        logger.error(f"Google Vision OCR failed: {e}")
        return ""


def _check_tesseract_kannada():
    """Check if Tesseract has Kannada language support installed."""
    try:
        available_langs = pytesseract.get_languages()
        if "kan" not in available_langs:
            logger.warning("Kannada language pack not found in Tesseract")
            logger.info("Available languages: " + ", ".join(available_langs))
            return False
        logger.info("Tesseract Kannada language pack found")
        return True
    except Exception as e:
        logger.error(f"Cannot check Tesseract languages: {e}")
        return False


def _process_ocr_text(text, is_legacy_detected=False, debug_mode=False):
    """Process OCR output with Kannada-specific corrections."""
    if not text.strip():
        return ""

    logger.debug(f"Processing OCR text: {len(text)} characters")

    try:
        if debug_mode:
            # DEBUG MODE: Skip legacy processing to see raw OCR quality
            logger.info("DEBUG MODE: Skipping legacy processing")
            processed = normalize_unicode(text)
        else:
            # Apply legacy conversion if needed
            processed = post_process_kannada_text(text, is_legacy=is_legacy_detected)

        # Validate output quality
        is_valid, issues = validate_kannada_output(processed)
        if not is_valid:
            logger.warning(f"OCR quality issues detected: {', '.join(issues)}")

        logger.debug(f"Processed text: {len(processed)} characters")
        return processed

    except Exception as e:
        logger.error(f"Text processing failed: {e}")
        return text  # Return original if processing fails


def _set_kannada_font(run):
    """Set appropriate font for Kannada text rendering."""
    try:
        run.font.name = "Noto Sans Kannada"
        # Fallback fonts for better compatibility
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Kannada")
        run._element.rPr.rFonts.set(qn("w:cs"), "Noto Sans Kannada")
    except Exception as e:
        logger.warning(f"Could not set Kannada font: {e}")


def ocr_pdf_to_word(
    input_pdf_path: str,
    output_docx_path: str,
    output_txt_path: str | None = None,
    *,
    language: str = "kan",
    title: str | None = None,
    author: str | None = None,
    use_google: bool = False,
    vision_page_limit: int | None = None,
    ocr_timeout: int = 30,
    debug_mode: bool = False,
) -> tuple[str, str]:
    """Perform OCR on a scanned PDF and output a Word document."""

    logger.info(f"Starting OCR conversion: {input_pdf_path}")
    logger.info(
        f"Settings: use_google={use_google}, vision_page_limit={vision_page_limit}, timeout={ocr_timeout}"
    )

    # Validate Tesseract setup for Kannada
    if not use_google and not _check_tesseract_kannada():
        raise RuntimeError(
            "Tesseract Kannada language pack not found. "
            "Install with: brew install tesseract-lang"
        )

    # Initialize document
    document = Document()

    # Set up output paths
    if output_txt_path is None:
        output_txt_path = os.path.splitext(output_docx_path)[0] + ".txt"

    # Set document metadata
    core = document.core_properties
    if title:
        core.title = title
    if author:
        core.author = author
    core.language = "kn-IN"  # Kannada locale

    full_text = ""
    total_pages = 0
    processed_pages = 0
    successful_pages = 0

    try:
        with TemporaryDirectory() as temp_dir:
            # Convert PDF to images with lower DPI for speed
            logger.info(f"Converting PDF to images")
            try:
                images = convert_from_path(
                    input_pdf_path,
                    output_folder=temp_dir,
                    fmt="png",
                    dpi=200,  # Reduced DPI for faster processing
                    thread_count=1,  # Single thread to avoid memory issues
                    first_page=1,
                    last_page=None,  # Process all pages
                )
            except Exception as e:
                logger.error(f"Failed to convert PDF to images: {e}")
                raise RuntimeError(f"Failed to convert PDF to images: {e}")

            total_pages = len(images)
            logger.info(f"Successfully converted {total_pages} pages to images")

            # Initialize Google Vision client if needed
            if use_google:
                try:
                    vision_client = vision.ImageAnnotatorClient()
                    logger.info("Google Vision client initialized")
                except Exception as e:
                    logger.error(f"Failed to initialize Google Vision: {e}")
                    use_google = False
                    logger.info("Falling back to Tesseract")

            for page_num, img in enumerate(images, start=1):
                logger.info(f"Processing page {page_num}/{total_pages}")
                processed_pages += 1
                page_text = ""

                try:
                    # Preprocess image for better OCR
                    processed_img = _preprocess_image_for_kannada(img)

                    # Determine OCR method
                    use_vision_for_page = use_google and (
                        vision_page_limit is None or page_num <= vision_page_limit
                    )

                    # Perform OCR with error handling
                    if use_vision_for_page:
                        logger.info(f"Using Google Vision for page {page_num}")
                        try:
                            page_text = _perform_google_vision_ocr(
                                processed_img, language, timeout=ocr_timeout
                            )
                        except TimeoutError:
                            logger.warning(
                                f"Google Vision timeout on page {page_num}, falling back to Tesseract"
                            )
                            page_text = _perform_tesseract_ocr(
                                processed_img, language, timeout=ocr_timeout // 2
                            )
                        except Exception as e:
                            logger.error(
                                f"Google Vision failed on page {page_num}, falling back to Tesseract: {e}"
                            )
                            page_text = _perform_tesseract_ocr(
                                processed_img, language, timeout=ocr_timeout // 2
                            )
                    else:
                        logger.info(f"Using Tesseract for page {page_num}")
                        page_text = _perform_tesseract_ocr(
                            processed_img, language, timeout=ocr_timeout
                        )

                    # DEBUG: Log raw OCR output
                    logger.info(
                        f"Raw OCR output for page {page_num}: '{page_text[:200]}...'"
                    )

                    # Detect legacy encoding
                    is_legacy = detect_legacy_encoding(page_text)
                    logger.info(
                        f"Legacy encoding detected on page {page_num}: {is_legacy}"
                    )

                    # DEBUG: Log before and after processing
                    logger.info(f"Before processing: '{page_text[:100]}...'")
                    cleaned_text = _process_ocr_text(
                        page_text, is_legacy, debug_mode=debug_mode
                    )
                    logger.info(f"After processing: '{cleaned_text[:100]}...'")

                    # DEBUG: Check if processing made it worse
                    if len(page_text) > 0 and len(cleaned_text) < len(page_text) * 0.1:
                        logger.warning(
                            f"Processing reduced text from {len(page_text)} to {len(cleaned_text)} chars - possible over-cleaning"
                        )

                    # Add page content to document
                    # Add page header
                    page_header = document.add_paragraph(f"ಪುಟ {page_num}")
                    page_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    page_header.style = "Heading 2"
                    for run in page_header.runs:
                        _set_kannada_font(run)

                    # Save and add page image (smaller size)
                    img_path = os.path.join(temp_dir, f"page_{page_num}.png")
                    # Resize image for document to save space
                    display_img = img.copy()
                    display_img.thumbnail((800, 1200), Image.Resampling.LANCZOS)
                    display_img.save(img_path, "PNG", optimize=True)

                    try:
                        document.add_picture(img_path, width=Inches(4))
                    except Exception as e:
                        logger.warning(f"Could not add image for page {page_num}: {e}")

                    # Add extracted text
                    if cleaned_text.strip():
                        text_para = document.add_paragraph(cleaned_text)
                        for run in text_para.runs:
                            _set_kannada_font(run)

                        full_text += cleaned_text + "\n\n"
                        successful_pages += 1
                        logger.info(
                            f"Successfully extracted text from page {page_num} ({len(cleaned_text)} chars)"
                        )
                    else:
                        no_text_para = document.add_paragraph(
                            "(ಈ ಪುಟದಲ್ಲಿ ಯಾವುದೇ ಪಠ್ಯ ಪತ್ತೆಯಾಗಿಲ್ಲ)"
                        )
                        for run in no_text_para.runs:
                            _set_kannada_font(run)
                            run.italic = True
                        logger.warning(f"No text extracted from page {page_num}")

                    # Add page break except for last page
                    if page_num < total_pages:
                        document.add_page_break()

                    # Clean up memory after each page
                    del processed_img, display_img
                    gc.collect()

                except Exception as e:
                    logger.error(f"Error processing page {page_num}: {e}")
                    # Add error message to document
                    error_para = document.add_paragraph(
                        f"ಪುಟ {page_num} ಸಂಸ್ಕರಣೆಯಲ್ಲಿ ದೋಷ: {str(e)}"
                    )
                    for run in error_para.runs:
                        _set_kannada_font(run)
                        run.italic = True

                    if page_num < total_pages:
                        document.add_page_break()

                    # Continue with next page
                    continue

            logger.info(
                f"OCR processing completed: {successful_pages}/{processed_pages} pages successful"
            )

    except Exception as e:
        logger.error(f"Error during OCR processing: {e}")
        # Add error to document instead of raising
        error_para = document.add_paragraph(f"ಸಂಸ್ಕರಣೆಯಲ್ಲಿ ದೋಷ: {str(e)}")
        for run in error_para.runs:
            _set_kannada_font(run)

    # Final text processing and validation
    if full_text.strip():
        full_text = normalize_unicode(full_text)

        # Validate final output
        is_valid, issues = validate_kannada_output(full_text)
        if not is_valid:
            logger.warning(f"Final output quality issues: {', '.join(issues)}")
    else:
        logger.error("No text was extracted from any page")
        full_text = "ಈ ದಾಖಲೆಯಿಂದ ಯಾವುದೇ ಪಠ್ಯವನ್ನು ಹೊರತೆಗೆಯಲಾಗಲಿಲ್ಲ."

    # Add summary to document
    summary_para = document.add_paragraph(
        f"\nಸಾರಾಂಶ: {successful_pages}/{total_pages} ಪುಟಗಳಿಂದ ಪಠ್ಯ ಹೊರತೆಗೆಯಲಾಗಿದೆ"
    )
    for run in summary_para.runs:
        _set_kannada_font(run)
        run.bold = True

    # Save Word document
    try:
        document.save(output_docx_path)
        logger.info(f"Word document saved: {output_docx_path}")
    except Exception as e:
        logger.error(f"Failed to save Word document: {e}")
        raise

    # Save text file
    try:
        with open(output_txt_path, "w", encoding="utf-8") as f:
            f.write(full_text)
        logger.info(f"Text file saved: {output_txt_path}")
    except Exception as e:
        logger.error(f"Failed to save text file: {e}")
        raise

    return output_docx_path, output_txt_path


def convert_scanned_pdf_to_word(*args, **kwargs):
    """Backward-compatible alias for :func:`ocr_pdf_to_word`."""
    return ocr_pdf_to_word(*args, **kwargs)
