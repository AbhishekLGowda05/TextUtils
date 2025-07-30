# -*- coding: utf-8 -*-
"""Simple PDF to Word conversion utilities using OCR.

This module provides a straightforward interface for converting
Kannada PDF files to Word documents. It includes an OCR based
converter that works for both digital PDFs encoded with legacy fonts
and scanned PDFs that only contain images.
"""

from __future__ import annotations

import io
import os
import tempfile
from typing import Optional

import fitz  # PyMuPDF
from PIL import Image
from docx import Document


__all__ = [
    "pdf_to_word_ocr_method",
    "pdf_to_word_with_ocr",
    "pdf_to_word_image_based",
    "setup_word_document_for_kannada",
    "setup_kannada_font_in_run",
    "pdf_to_word",
]


def pdf_to_word_ocr_method(file) -> str:
    """Convert a PDF file handle to a Word document using OCR.

    The returned value is the path to the generated Word file.
    """
    try:
        os.makedirs("static/downloads", exist_ok=True)
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        file.save(temp_pdf.name)
        temp_pdf.close()

        output_path = os.path.join("static/downloads", "converted.docx")

        try:
            if pdf_to_word_with_ocr(temp_pdf.name, output_path):
                os.unlink(temp_pdf.name)
                return output_path
        except Exception:
            pass

        if pdf_to_word_image_based(temp_pdf.name, output_path):
            os.unlink(temp_pdf.name)
            return output_path
        os.unlink(temp_pdf.name)
        raise RuntimeError("PDF conversion failed")
    except Exception as exc:  # pragma: no cover - unexpected errors
        raise Exception(f"PDF to Word conversion failed: {exc}") from exc


def pdf_to_word_with_ocr(pdf_path: str, output_path: str) -> bool:
    """Convert PDF to Word using OCR (Tesseract)."""
    try:
        import pytesseract
    except Exception:
        print("pytesseract not available")
        return False

    try:
        pytesseract.get_tesseract_version()
    except Exception:
        print("Tesseract OCR not installed")
        return False

    pdf_doc = fitz.open(pdf_path)
    word_doc = Document()
    setup_word_document_for_kannada(word_doc)

    try:
        available = pytesseract.get_languages()  # type: ignore[arg-type]
        lang = "kan+eng" if "kan" in available else "eng"
    except Exception:
        lang = "eng"

    for page_num in range(len(pdf_doc)):
        page = pdf_doc.load_page(page_num)
        pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0))
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        try:
            text = pytesseract.image_to_string(
                img,
                lang=lang,
                config="--oem 3 --psm 6",
            )
        except Exception as err:
            print(f"OCR failed for page {page_num+1}: {err}")
            text = ""

        if text.strip():
            if page_num > 0:
                word_doc.add_page_break()
            for para in text.split("\n\n"):
                para = para.strip()
                if para:
                    run = word_doc.add_paragraph().add_run(para)
                    setup_kannada_font_in_run(run)

    word_doc.save(output_path)
    pdf_doc.close()
    return True


def pdf_to_word_image_based(pdf_path: str, output_path: str) -> bool:
    """Fallback conversion that embeds each PDF page as an image."""
    try:
        pdf_doc = fitz.open(pdf_path)
        word_doc = Document()
        setup_word_document_for_kannada(word_doc)

        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
            img = Image.open(io.BytesIO(pix.tobytes("png")))

            temp_img = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            img.save(temp_img.name, format="PNG")
            temp_img.close()

            if page_num > 0:
                word_doc.add_page_break()
            run = word_doc.add_paragraph().add_run()
            from docx.shared import Inches

            run.add_picture(temp_img.name, width=Inches(6))
            os.unlink(temp_img.name)

        word_doc.save(output_path)
        pdf_doc.close()
        return True
    except Exception as err:
        print(f"Image-based conversion error: {err}")
        return False


def setup_word_document_for_kannada(word_doc: Document) -> None:
    """Configure default styles for Kannada text."""
    from docx.oxml.ns import qn

    styles = word_doc.styles
    default_style = styles["Normal"]
    font = default_style.font
    font.name = "Nirmala UI"

    rpr = default_style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(key), "Nirmala UI")


def setup_kannada_font_in_run(run) -> None:
    """Apply Kannada font settings to a run."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = "Nirmala UI"
    run.font.size = Pt(12)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(key), "Nirmala UI")
    lang = rpr.get_or_add_lang()
    lang.set(qn("w:bidi"), "kn-IN")


def pdf_to_word(file) -> str:
    """Public helper that delegates to :func:`pdf_to_word_ocr_method`."""
    return pdf_to_word_ocr_method(file)
