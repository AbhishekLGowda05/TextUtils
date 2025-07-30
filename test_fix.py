#!/usr/bin/env python3
"""
Quick test script to verify the fixes for both digital and scanned PDFs.
Run this to test the improvements without going through the web interface.
"""

import os
import sys
import logging
from pathlib import Path

# Add modules to path
sys.path.append(os.path.join(os.path.dirname(__file__), 'modules'))

from modules.pdf_to_word import convert_pdf_to_word
from modules.ocr_to_word import convert_scanned_pdf_to_word

# Set up logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

def test_digital_pdf():
    """Test digital PDF processing with improved legacy detection."""
    print("\n" + "="*50)
    print("TESTING DIGITAL PDF PROCESSING")
    print("="*50)
    
    # Look for any PDF in uploads directory
    uploads_dir = Path("uploads")
    pdf_files = list(uploads_dir.glob("*.pdf"))
    
    if not pdf_files:
        print("❌ No PDF files found in uploads directory")
        return
    
    pdf_file = pdf_files[0]
    print(f"📄 Testing with: {pdf_file.name}")
    
    try:
        # Test with debug mode to bypass legacy processing
        result_path = convert_pdf_to_word(
            str(pdf_file), 
            "static/converted", 
            debug_mode=True  # This should prevent legacy corruption
        )
        
        if result_path and os.path.exists(result_path):
            print(f"✅ Digital PDF conversion successful: {result_path}")
            
            # Read a sample of the output to check quality
            try:
                from docx import Document
                doc = Document(result_path)
                sample_text = ""
                for para in doc.paragraphs[:3]:  # First 3 paragraphs
                    if para.text.strip():
                        sample_text += para.text[:100] + "\n"
                
                print(f"📝 Sample output text:\n{sample_text}")
                
                # Check for corruption signs
                corruption_signs = ["ೇೆ", "ವಿ ಪೄೆ", "K ಯುವವ", "ರರರ"]
                has_corruption = any(sign in sample_text for sign in corruption_signs)
                
                if has_corruption:
                    print("⚠️  WARNING: Possible corruption detected in output")
                else:
                    print("✅ Text appears clean (no obvious corruption)")
                    
            except Exception as e:
                print(f"⚠️  Could not read output file: {e}")
        else:
            print("❌ Digital PDF conversion failed")
            
    except Exception as e:
        print(f"❌ Digital PDF test failed: {e}")

def test_scanned_pdf():
    """Test scanned PDF processing with improved OCR."""
    print("\n" + "="*50)
    print("TESTING SCANNED PDF PROCESSING")
    print("="*50)
    
    # Look for any PDF in uploads directory
    uploads_dir = Path("uploads")
    pdf_files = list(uploads_dir.glob("*.pdf"))
    
    if not pdf_files:
        print("❌ No PDF files found in uploads directory")
        return
    
    pdf_file = pdf_files[0]
    print(f"📄 Testing with: {pdf_file.name}")
    
    try:
        # Test with reduced page limit for quick testing
        result_path = convert_scanned_pdf_to_word(
            str(pdf_file), 
            "static/converted",
            max_pages=2,  # Only process first 2 pages for quick test
            google_vision_pages=2,  # Use Google Vision for better quality
            debug_mode=True  # Enable debug logging
        )
        
        if result_path and os.path.exists(result_path):
            print(f"✅ Scanned PDF conversion successful: {result_path}")
            
            # Read a sample of the output to check quality
            try:
                from docx import Document
                doc = Document(result_path)
                sample_text = ""
                for para in doc.paragraphs[:3]:  # First 3 paragraphs
                    if para.text.strip():
                        sample_text += para.text[:100] + "\n"
                
                print(f"📝 Sample output text:\n{sample_text}")
                
                if sample_text.strip():
                    print("✅ OCR produced text output")
                else:
                    print("⚠️  WARNING: OCR produced empty output")
                    
            except Exception as e:
                print(f"⚠️  Could not read output file: {e}")
        else:
            print("❌ Scanned PDF conversion failed")
            
    except Exception as e:
        print(f"❌ Scanned PDF test failed: {e}")

def main():
    """Run all tests."""
    print("🧪 PDF CONVERSION TESTING SUITE")
    print("Testing the recent fixes...")
    
    # Test digital PDF processing
    test_digital_pdf()
    
    # Test scanned PDF processing  
    test_scanned_pdf()
    
    print("\n" + "="*50)
    print("✅ TESTING COMPLETE")
    print("Check the results above to see if the fixes are working.")
    print("If issues persist, check the debug logs for more details.")
    print("="*50)

if __name__ == "__main__":
    main()
